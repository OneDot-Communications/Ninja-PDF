import sys
import os
import subprocess
import traceback
import tempfile

# Create a dummy docx file for testing
from docx import Document
try:
    doc = Document()
    doc.add_paragraph("Test Document for PDF Conversion")
    test_docx = os.path.abspath("test_debug.docx")
    doc.save(test_docx)
    print(f"Created test file: {test_docx}")
except ImportError:
    # If python-docx not installed, just write a text file and rename (docx2pdf might fail validation but we test execution)
    # Actually docx2pdf opens it in Word, so it must be valid.
    # Let's try to install python-docx if missing? using pip
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test Document for PDF Conversion")
    test_docx = os.path.abspath("test_debug.docx")
    doc.save(test_docx)
    print(f"Created test file: {test_docx}")

output_pdf = os.path.abspath("test_debug.pdf")

print(f"Python: {sys.executable}")
print("Attempting conversion via docx2pdf module...")

try:
    cmd = [sys.executable, '-m', 'docx2pdf', test_docx, output_pdf]
    print(f"Command: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    print(f"Return Code: {result.returncode}")
    print(f"STDOUT: {result.stdout}")
    print(f"STDERR: {result.stderr}")
    
    if result.returncode == 0 and os.path.exists(output_pdf):
        print("SUCCESS: Conversion worked.")
    else:
        print("FAILURE: Conversion failed.")

except Exception as e:
    print(f"EXCEPTION: {e}")
    traceback.print_exc()
